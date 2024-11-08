import fitz  # PyMuPDF for PDF extraction  
from openai import AzureOpenAI  
from dotenv import load_dotenv  
import os  
import re  # For parsing structured output  z
import json  # For JSON handling  
import pandas as pd  
import streamlit as st  
import docx  
from io import BytesIO  
from azure.ai.formrecognizer import DocumentAnalysisClient  
from azure.core.credentials import AzureKeyCredential  
from azure.core.exceptions import HttpResponseError  
from docx2pdf import convert  
import pypandoc  
from PyPDF2 import PdfMerger  
import tempfile
import nltk  
from nltk.tokenize import word_tokenize  
from nltk.corpus import stopwords  
from nltk.stem import WordNetLemmatizer  
import string
import logging
from pydantic import BaseModel, ValidationError 
import time  
import random  
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from azure.storage.blob import BlobServiceClient  
  
# Define your connection string and container name  
connection_string = "DefaultEndpointsProtocol=https;AccountName=patent;AccountKey=F8W2RDEJMD5kJ3elvulDyz6XJlYBWp4K3WY4IHhTHKfW+tG/HKTVKmzZS5y0J6GgWFp1uBKImwvs+ASthawnOA==;EndpointSuffix=core.windows.net"  
container_name = "patent"  
  
# Initialize the BlobServiceClient  
blob_service_client = BlobServiceClient.from_connection_string(connection_string)  
  
# Create the container if it doesn't exist  
container_client = blob_service_client.get_container_client(container_name)  
try:  
    container_client.create_container()  
except Exception as e:  
    print(f"Container already exists or could not be created: {e}")  

def upload_file_to_blob(file_path, blob_name, folder_name=None):  
    try:  
        # Sanitize the folder name to ensure it doesn't contain any directory separators  
        if folder_name:  
            folder_name = folder_name.replace('/', '').replace('\\', '')  
  
        # Determine the full blob name with folder structure  
        full_blob_name = f"{folder_name}/{blob_name}" if folder_name else blob_name  
  
        # Get a blob client  
        blob_client = container_client.get_blob_client(full_blob_name)  
  
        # Upload the file  
        with open(file_path, "rb") as data:  
            blob_client.upload_blob(data, overwrite=True)  
  
        print(f"File {file_path} uploaded to blob {full_blob_name}.")  
    except Exception as e:  
        print(f"Error uploading file to blob: {e}")  

# Configure logging  
logging.basicConfig(  
    level=logging.INFO,  # Change to DEBUG for more detailed logs  
    format='%(asctime)s - %(levelname)s - %(message)s',  
    handlers=[logging.FileHandler("app.log"), logging.StreamHandler()]  
)  


# Make sure to download the necessary NLTK data  
nltk.download('punkt')  
nltk.download('stopwords')  
nltk.download('wordnet')  
nltk.download('omw-1.4') 
nltk.download('punkt_tab')
# Load environment variables from .env file  
load_dotenv()  
# Initialize global variables  
domain_subject_matter = "default domain"  
experience_expertise_qualifications = "default qualifications"  
style_tone_voice = "default style"

 
# Set up Azure OpenAI API credentials from .env  
client = AzureOpenAI(  
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),  # Pull from environment  
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),  # Pull from environment  
    api_version=os.getenv("OPENAI_API_VERSION"),  # Pull from environment  
)  
  
# Azure Form Recognizer setup  
form_recognizer_endpoint = os.getenv("FORM_RECOGNIZER_ENDPOINT")  
form_recognizer_api_key = os.getenv("FORM_RECOGNIZER_API_KEY") 
# Define a model for conflict results  
class DomainExpertise(BaseModel):  
    domain_subject_matter: str  
    experience_expertise_qualifications: str  
    style_tone_voice: str   

# Define the Pydantic model for extracted details  
class FoundationalClaimDetails(BaseModel):  
    foundational_claim_details: list[dict]  

# Preprocessing function  
def process_text(text):
 logging.info("Started processing text.")   
 try:
    """Process the extracted text by tokenizing, removing stop words, punctuation, extra whitespace, and lemmatizing."""  
    stop_words = set(stopwords.words('english'))  
    lemmatizer = WordNetLemmatizer()  
      
    words = word_tokenize(text)  
      
    processed_words = []  
    for word in words:  
        word_lower = word.lower()  # Convert to lowercase  
        if word_lower not in stop_words and word_lower not in string.punctuation:  
            lemmatized_word = lemmatizer.lemmatize(word_lower)  
            processed_words.append(lemmatized_word)  
      
    processed_text = ' '.join(processed_words)  
    processed_text = ' '.join(processed_text.split())  
    logging.info("Completed processing text.")    
    return processed_text  
 except Exception as e:  
        logging.error(f"Error processing text: {e}")  
        raise    
def process_text_with_page_numbers(file_content):  
    logging.info("Started processing text with page numbers.")  
    try:  
        stop_words = set(stopwords.words('english'))  
        lemmatizer = WordNetLemmatizer()  
        processed_pages = []  
  
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key)  
        )  
  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=file_content  
        )  
        result = poller.result()  
  
        for page in result.pages:  
            page_number = page.page_number  
            text = "\n".join([line.content for line in page.lines])  
  
            words = word_tokenize(text)  
            processed_words = [  
                lemmatizer.lemmatize(word.lower())  
                for word in words  
                if word.lower() not in stop_words and word.lower() not in string.punctuation  
            ]  
            processed_text = ' '.join(processed_words)  
            processed_text = ' '.join(processed_text.split())  
  
            processed_text_with_page = f"{processed_text} [Page {page_number}]"  
            processed_pages.append(processed_text_with_page)  
  
        logging.info("Completed processing text with page numbers.")  
        return processed_pages  
    except Exception as e:  
        logging.error(f"Error processing text: {e}")  
        raise  
  
def extract_and_process_text_from_pdf(uploaded_pdf_path):  
    try:  
        with open(uploaded_pdf_path, "rb") as f:  
            file_content = f.read()  
            processed_pages = process_text_with_page_numbers(file_content)  
            return processed_pages  # Return processed pages with page numbers  
    except Exception as e:  
        logging.error(f"Failed to process the file: {e}")  
        return None  
  
  
def extract_text_from_docx(uploaded_docx): 
 logging.info(f"Extracting text from DOCX: {uploaded_docx.name}") 
 try: 
    """Extract text from a DOCX file."""  
    doc = docx.Document(uploaded_docx)  
    full_text = []  
    for para in doc.paragraphs:  
        full_text.append(para.text)  
    return "\n".join(full_text)  
 except Exception as e:  
        logging.error(f"Error extracting text from DOCX: {e}")  
        raise  
  
def determine_domain_expertise(action_document_text):  
    """Analyze the action document to determine the required domain expertise, experience, and analysis style."""  
    global domain_subject_matter, experience_expertise_qualifications, style_tone_voice  
  
    prompt = f"""Analyze the following action document text and determine the domain expertise required to analyze this document: {action_document_text}  
    Step 1: Identify the subject matter and domain expertise required to understand this document and the cited documents in depth.  
    Step 2: Determine the experience, expertise, and educational qualifications required to handle this document and the cited documents in depth.  
    Step 3: Describe the style, tone, and voice required to analyze these kinds of documents.  
    NOTE: Each answer needs to be detailed.  
    Step 4: Provide the response in the following JSON format:  
    {{  
        "domain_subject_matter": "Detailed description of the domain subject matter",  
        "experience_expertise_qualifications": "Detailed description of the experience, expertise, and educational qualifications required",  
        "style_tone_voice": "Detailed description of the style, tone, and voice required"  
    }}  
    """  
  
    messages = [  
        {"role": "system", "content": "You are an AI assistant that can analyze the following action document text and determine the domain, expertise, and subject matter required to analyze this document."},  
        {"role": "user", "content": prompt}  
    ]  
  
    try:  
        # Call OpenAI API for domain expertise determination  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
  
        # Extract the content from the response  
        raw_content = response.choices[0].message.content.strip()  
  
        # Print the raw response for debugging  
        print("Raw API Response:")  
        print(raw_content)  
  
        # Use regex to find JSON content  
        json_match = re.search(r'{.*}', raw_content, re.DOTALL)  
        if not json_match:  
            print("No JSON content found in the response.")  
            return (None, None, None)  
  
        cleaned_content = json_match.group(0)  
  
        # Check for code block markers and remove them  
        if cleaned_content.startswith("```json"):  
            cleaned_content = cleaned_content[7:-3].strip()  
        elif cleaned_content.startswith("```"):  
            cleaned_content = cleaned_content[3:-3].strip()  
  
        # Print the cleaned content for debugging  
        print("Cleaned API Response:")  
        print(cleaned_content)  
  
        # Validate and parse using Pydantic  
        try:  
            # Parse cleaned content as JSON to ensure it's valid  
            json_data = json.loads(cleaned_content)  
  
            # Validate with Pydantic model  
            expertise_data = DomainExpertise(**json_data)  
            domain_subject_matter = expertise_data.domain_subject_matter  
            experience_expertise_qualifications = expertise_data.experience_expertise_qualifications  
            style_tone_voice = expertise_data.style_tone_voice  
            return (domain_subject_matter, experience_expertise_qualifications, style_tone_voice)  
        except (ValidationError, json.JSONDecodeError) as e:  
            print(f"Validation or JSON error: {str(e)}")  
            return (None, None, None)  
  
    except Exception as e:  
        print(f"Error during domain expertise determination: {str(e)}")  
        return (None, None, None)  
  
    except Exception as e:  
        st.error(f"Error during domain expertise determination: {str(e)}")  
        return (None, None, None) 
    
# Define the Pydantic model for validation  
class ConflictResults(BaseModel):  
    foundational_claim: str  
    documents_referenced: list  
    page_no: list  
    figures: list  
    text: str  
  
def escape_curly_braces(text):  
    """Escape curly braces in the given text."""  
    return text.replace("{", "{{").replace("}", "}}")  
  
def generate_system_content(domain, expertise, style):  
    """Generate the system content based on domain, expertise, and style."""  
    return f"""  
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:  
    1. {domain}  
    2. Patent Law Proficiency:  
        a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.  
        b. Knowledgeable about the structure and requirements of patent applications.  
        c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
    3. {expertise}  
    4. Capability to Propose Amendments:  
        a. Experienced in responding to examiners’ assertions or rejections of claims.  
        b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
        c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.  
    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:  
    a. A thorough evaluation of the technical details and functionalities described in the patent application.  
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.  
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.  
    d. Feedback on the strengths and potential areas for improvement in the document.  
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.  
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.  
    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.  
    """  
  
def generate_user_prompt(escaped_text):  
    """Generate the user prompt for conflict analysis."""  
    return f"""  
    Analyze the following action document text and extract the foundational claim:  
    {escaped_text}  
    Step 1: Extract the key claims from the document and name it as 'Key_claims'.  
    Step 2: From the 'Key_claims' extract the foundational claim with its number and store it in a variable called "foundational_claim" (Note: method claims and system claims are not considered independent claims and only one claim can be the foundational claim).  
    Step 3: From the foundational claim, extract the information under U.S.C 102 and/or 103.  
    Step 4: Extract all referenced documents under U.S.C. 102 and/or 103 mentioned in the action document specified only in the "foundational_claim".  
    Step 5: For each referenced document, create a variable that stores the document name.  
    Step 6: If the foundational claim refers to the referenced documents, extract the entire technical content with its specified paragraph location and image reference. Map the claim with the conflicting document name.  
    Step 7: Do not extract any referenced document data that is not related to the foundational claim.  
    NOTE: Extract in English.  
    NOTE: Give the documents referenced with their publication numbers EG. Deker US...  
    Step 8: Return the output as a JSON object with the following structure:  
    {{  
        "foundational_claim": "text",  
        "documents_referenced": ["doc1", "doc2", ...],  
        "page_no:" ["page 19","page 21"]  
        "figures": ["fig1", "fig2", ...],  
        "text": "detailed text"  
    }}  
    """  
  
def call_api_with_retries(messages):  
    """Call the API with retries."""  
    max_retries = 3  
    retry_delay = 2  # seconds  
    for attempt in range(max_retries):  
        try:  
            response = client.chat.completions.create(  
                model="GPT-4-Omni", messages=messages, temperature=0.2  
            )  
            return response  
        except Exception as e:  
            logging.error(f"API call failed on attempt {attempt + 1}: {str(e)}")  
            if attempt < max_retries - 1:  
                time.sleep(retry_delay)  
            else:  
                raise  
  
def extract_json_response(content):  
    """Extract JSON from the response content."""  
    start_index = content.find("```json")  
    if start_index != -1:  
        end_index = content.find("```", start_index + 7)  
        if end_index != -1:  
            return content[start_index + 7:end_index].strip()  
        else:  
            return content[start_index + 7:].strip()  
    return None  
  
def validate_and_parse_json(json_string):  
    """Validate and parse the JSON string."""  
    if json_string:  
        try:  
            json_data = json.loads(json_string)  
            conflict_results = ConflictResults(**json_data)  
            return conflict_results.dict()  
        except json.JSONDecodeError as e:  
            logging.error(f"JSON decoding error: {str(e)}")  
            logging.error(f"Content causing error: {json_string}")  
        except ValidationError as e:  
            logging.error(f"Validation error: {str(e)}")  
            logging.error(f"Content causing error: {json_string}")  
    else:  
        logging.error("No JSON content extracted.")  
    return None  
  
def check_for_conflicts(action_document_text, domain, expertise, style):  
    """Main function to check for conflicts in the action document."""  
    escaped_text = escape_curly_braces(action_document_text)  
    content = generate_system_content(domain, expertise, style)  
    prompt = generate_user_prompt(escaped_text)  
  
    messages = [  
        {"role": "system", "content": content},  
        {"role": "user", "content": prompt},  
    ]  
  
    try:  
        response = call_api_with_retries(messages)  
        content = response.choices[0].message.content.strip()  
        json_string = extract_json_response(content)  
        return validate_and_parse_json(json_string)  
    except Exception as e:  
        logging.error(f"Error during conflict checking: {str(e)}")  
        return None 


# Define the Pydantic model for validation  
class FigureAnalysisResults(BaseModel):  
    figures_analysis: list  
    extracted_paragraphs: list  
  
def generate_system_content(domain, expertise, style):  
    """Generate system content based on domain, expertise, and style."""  
    return f"""  
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:  
    1. {domain}  
    2. Patent Law Proficiency:  
        a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.  
        b. Knowledgeable about the structure and requirements of patent applications.  
        c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
    3. {expertise}  
    4. Capability to Propose Amendments:  
        a. Experienced in responding to examiners’ assertions or rejections of claims.  
        b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
        c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.  
    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:  
    a. A thorough evaluation of the technical details and functionalities described in the patent application.  
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.  
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.  
    d. Feedback on the strengths and potential areas for improvement in the document.  
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.  
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.  
    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.  
    """  
  
def generate_figure_analysis_prompt(fig_details, text_details, ref_documents_texts):  
    """Generate prompt for figure analysis."""  
    return f"""  
    Analyze the figures and technical text from the referenced document in relation to the foundational claim.  
    Instructions:  
    1. Identify Figures:  
        - For each figure referenced in the foundational claim, extract the following:  
            - **Figure Number and Title:** Provide the figure number and its title.  
            - **Technical Details:** Extract all technical details related to the figure as mentioned in the text. Ensure no technical detail is missed.  
            - **Importance:** Explain the importance of the figure in relation to the foundational claim. Describe how it supports, illustrates, or contradicts the claim.  
    2. Extract Text from Paragraphs:  
        - From the paragraphs cited in the foundational claim, extract the relevant text as in the document uploaded and store it in a separate variable.  
    3. Workflow for Cases with Images:  
        - If figures are present in the referenced document:  
            - Follow the steps outlined above to extract figure details and technical information.  
            - Ensure that any interpretations of the figures include specific references to the data or concepts depicted.  
    4. Workflow for Cases without Images:  
        - If no figures are present:  
            - Focus on extracting and analyzing the text from the referenced document.  
            - Identify and highlight key technical details and concepts that are essential to understanding the foundational claim.  
    Input Details:  
    Figures: {json.dumps(fig_details, indent=2)}  
    Text: {text_details}  
    Referenced Document Texts: {json.dumps(ref_documents_texts, indent=2)}  
    Response format:  
    {{  
        "figures_analysis": [  
            {{  
                "figure_number": "Figure 1",  
                "title": "Title of Figure 1",  
                "technical_details": "Detailed technical description",  
                "importance": "Explanation of importance"  
            }},  
            ...  
        ],  
        "extracted_paragraphs": [  
            "Paragraph text 1",  
            ...  
        ]  
    }}  
    """  
  
def call_llm_api(messages):  
    """Call the LLM API and handle the response."""  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.2  
        )  
        return response.choices[0].message.content.strip()  
    except Exception as e:  
        print(f"Unexpected error: {e}")  
        return None  
  
def parse_and_validate_json(analysis_output):  
    """Parse and validate the JSON output from LLM."""  
    if analysis_output.startswith("```json"):  
        analysis_output = analysis_output[7:-3].strip()  
    elif analysis_output.startswith("```"):  
        analysis_output = analysis_output[3:-3].strip()  
  
    if analysis_output:  
        try:  
            json_data = json.loads(analysis_output)  
            figure_analysis_results = FigureAnalysisResults(**json_data)  
            return figure_analysis_results.dict()  
        except json.JSONDecodeError as e:  
            print(f"JSON decoding error during validation: {e}")  
            print(f"Analysis output content causing error: {analysis_output}")  
        except ValidationError as e:  
            print(f"Validation error: {e.json()}")  
            print(f"Analysis output content causing error: {analysis_output}")  
    else:  
        print("No content received from OpenAI API.")  
    return None  
  
def extract_figures_and_text(conflict_results, ref_documents_texts, domain, expertise, style):  
    """Main function to extract figures and related technical text."""  
    fig_details = conflict_results.get("figures", [])  
    text_details = conflict_results.get("text", "")  
  
    content = generate_system_content(domain, expertise, style)  
    figure_analysis_prompt = generate_figure_analysis_prompt(fig_details, text_details, ref_documents_texts)  
  
    messages = [  
        {"role": "system", "content": content},  
        {"role": "user", "content": figure_analysis_prompt},  
    ]  
  
    analysis_output = call_llm_api(messages)  
    return parse_and_validate_json(analysis_output) 

def generate_system_content(domain, expertise, style):  
    return f"""  
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:  
    1. {domain}  
    2. Patent Law Proficiency:  
        a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.  
        b. Knowledgeable about the structure and requirements of patent applications.  
        c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
    3. {expertise}  
    4. Capability to Propose Amendments:  
        a. Experienced in responding to examiners’ assertions or rejections of claims.  
        b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
        c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.  
    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:  
    a. A thorough evaluation of the technical details and functionalities described in the patent application.  
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.  
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.  
    d. Feedback on the strengths and potential areas for improvement in the document.  
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.  
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.  
    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.  
    """  
  
def generate_user_prompts(filed_application_text, foundational_claim):  
    return f"""  
    Analyze the following filed application text and extract details related to the foundational claim.  
    Filed Application Text: {filed_application_text}  
    Foundational Claim: {json.dumps(foundational_claim, indent=2)}  
    Instructions:  
    1. Identify and extract all technical details from the filed application that relate to the foundational claim.  
    2. Ensure that any extracted details include specific references to the paragraphs or sections in the filed application where they are found. NOTE: Extract in English.  
    3. Return the extracted details in the following JSON format:  
    JSON format:  
    {{  
        "foundational_claim_details": [  
            {{  
                "paragraph_number": "Paragraph 1",  
                "text": "Technical text relating to the foundational claim"  
            }},  
            ...  
        ]  
    }}  
    """  
  
def extract_json_contents(response_content):  
    start_index = response_content.find("```json")  
    if start_index != -1:  
        end_index = response_content.find("```", start_index + 7)  
        if end_index != -1:  
            return response_content[start_index + 7:end_index].strip()  
        return response_content[start_index + 7:].strip()  
    return response_content.strip()  
  
def validate_and_parse_jsons(json_string):  
    try:  
        parsed_json = json.loads(json_string)  
        details = FoundationalClaimDetails(**parsed_json)  
        return details.dict()  
    except json.JSONDecodeError as e:  
        print(f"JSON decoding error: {e}")  
        print(f"Raw response: {json_string}")  
    except ValidationError as e:  
        print(f"Validation error: {e.json()}")  
        print(f"Raw response: {json_string}")  
    return None  
  
def extract_details_from_filed_application(filed_application_text, foundational_claim, domain, expertise, style):  
    """Extract details from the filed application related to the foundational claim."""  
    content = generate_system_content(domain, expertise, style)  
    prompt = generate_user_prompts(filed_application_text, foundational_claim)  
  
    messages = [  
        {"role": "system", "content": content},  
        {"role": "user", "content": prompt},  
    ]  
  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.2  
        )  
  
        content = response.choices[0].message.content.strip()  
        print(f"Raw response: {content}")  
  
        json_string = extract_json_contents(content)  
        if json_string:  
            return validate_and_parse_jsons(json_string)  
        else:  
            print("No JSON content extracted.")  
            return None  
    except Exception as e:  
        print(f"Error extracting details from filed application: {e}")  
        return None 


  
# Function to extract details from pending claims and modify the filed application details  
def extract_and_modify_filed_application(filed_application_details, pending_claims_text, domain, expertise, style):  
    """  
    Extract details from the pending claims and modify the filed application details.  
    """
    content = f"""
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """  
    
    global domain_subject_matter, experience_expertise_qualifications, style_tone_voice   
    prompt = f"""  
    Analyze the following pending claims text and modify the filed application details accordingly.  
    Pending Claims Text: {pending_claims_text}  
    Filed Application Details: {json.dumps(filed_application_details, indent=2)}  
    Instructions:  
    1. Identify and extract all technical details from the pending claims that relate to the foundational claim.  
    2. Modify the filed application details based on the extracted details from the pending claims.  
    3. Ensure that any modifications include specific references to the paragraphs or sections in the pending claims where they are found.NOTE:Extract in English.  
    4. Return the modified filed application details in the following JSON format:  
    {{  
        "modified_filed_application_details": [  
            {{  
                "paragraph_number": "Paragraph 1",  
                "text": "Modified detailed text based on pending claims"  
            }},  
            ...  
        ]  
    }}  
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": content,
        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    # Call OpenAI API for extracting and modifying filed application details  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.2  
        )  
          
        # Extract the response content  
        content = response.choices[0].message.content.strip()  
  
        # Locate the JSON within triple backticks  
        start_index = content.find("```json")  
        if start_index != -1:  
            end_index = content.find("```", start_index + 7)  
            if end_index != -1:  
                json_string = content[start_index + 7:end_index].strip()  
            else:  
                json_string = content[start_index + 7:].strip()  
        else:  
            # If no JSON block is found, treat the entire content as potential JSON  
            json_string = content  
  
        # Print raw response for debugging  
        print(f"Raw response: {content}")  
  
        # Validate JSON structure  
        if json_string:  
            try:  
                # Parse the JSON to ensure it's valid  
                parsed_json = json.loads(json_string)  
                # Validate with Pydantic model  
                details = FoundationalClaimDetails(**parsed_json)  
                return details.dict()  
            except json.JSONDecodeError as e:  
                print(f"JSON decoding error: {e}")  
                print(f"Raw response: {json_string}")  
                return None  
            except ValidationError as e:  
                print(f"Validation error: {e.json()}")  
                print(f"Raw response: {json_string}")  
                return None  
        else:  
            print("No JSON content extracted.")  
            return None  
    except Exception as e:  
        print(f"Error extracting details from filed application: {e}")  
        return None 
  
 
def generate_content(domain, expertise, style):  
    """Generate system content for the analysis."""  
    return f"""  
    You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:  
    1. {domain}  
    2. Patent Law Proficiency:  
        a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.  
        b. Knowledgeable about the structure and requirements of patent applications.  
        c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
    3. {expertise}  
    4. Capability to Propose Amendments:  
        a. Experienced in responding to examiners’ assertions or rejections of claims.  
        b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).  
        c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.  
  
    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:  
    a. A thorough evaluation of the technical details and functionalities described in the patent application.  
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.  
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.  
    d. Feedback on the strengths and potential areas for improvement in the document.  
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.  
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.  
  
    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.  
    """  
  
def generate_few_shot_examples():  
    """Generate few-shot examples for analysis."""  
    few_shot_example = """  
    **Example Amendment and Argument:**  
    **Amendment 1: Enhanced Communication Protocol**  
    **Original Claim Language:**  
    "A communication system comprising a transmitter and receiver."  
    **Proposed Amended Language:**  
    "A communication system comprising a transmitter and receiver, wherein the transmitter is configured to utilize an adaptive frequency hopping protocol to dynamically adjust communication channels based on interference levels."  
    **Derivation and Reasoning:**  
    - **Source Reference**: Derived from Paragraphs [0040]-[0045] and Figures 4A-4D of the application.  
    - **Reasoning**: The amendment specifies the use of an "adaptive frequency hopping protocol" and includes dynamic adjustments based on interference levels, adding specificity and distinguishing over prior art that lacks adaptive frequency hopping.  
    **Supporting Arguments:**  
    - **Novelty**: The cited reference does not disclose a communication system utilizing an adaptive frequency hopping protocol that adjusts based on interference levels.  
    - **Non-Obviousness**: Combining a communication system with an adaptive frequency hopping protocol introduces an unexpected technical advantage by improving communication reliability and reducing interference, which is not suggested or rendered obvious by the prior art.  
    - **Technical Advantages**: Enhances communication reliability and reduces interference, as detailed in Paragraph [0046] of the application.  
    - **Addressing Examiner's Rejection**: The prior art only teaches static frequency selection methods, thus the amendment overcomes the rejection by introducing adaptive frequency hopping functionality not suggested in the cited reference.  
    """  
  
    text_a = """  
    **Example Analysis**  
    **Key Features of Independent Claim 1**  
    • **Multiparameter Leadset:** Configured to interface with a monitoring device for monitoring multiple health indicators of a patient.  
    • **Single Patient Plug:** Having a plurality of monitoring contacts.  
  
    **Key Features of Cited Reference(Naylor):**  
    • **Multiparameter Leadset (Naylor)**: Depicted in Figure 2 (page 19), comprising a temperature sensor, non-invasive pulse oximetry sensor, and EKG sensor.  
    • **Junction Box**: Connects to a patient monitor via a common output cable, with receptacles for each sensor plug (Figure 2: junction box 226).  
  
    **Examiner’s Analysis:**  
    The examiner rejected the application based on U.S.C 103 (Obviousness), asserting that the claimed features are either disclosed or obvious in light of the Naylor reference combined with Morley for the interconnection feature. The examiner interprets the cited reference as teaching or suggesting all elements of the foundational claim, including the use of a multiparameter leadset with a single patient plug and various patient leads for different health indicators. The interconnection feature is deemed obvious for better wire management.  
  
    **Novelty Analysis (U.S.C 102 - Lack of Novelty):**  
    Comparing the foundational claim with the Naylor reference:  
    • **Multiparameter Leadset**: Both the foundational claim and Naylor describe a multiparameter leadset.  
    • **Single Patient Plug**: Naylor's junction box 226 serves a similar function.  
  
    **Non-Obviousness Analysis (U.S.C 103 - Obviousness):**  
    The foundational claim may be considered obvious in light of Naylor combined with Morley:  
    • The interconnection feature, while not explicitly taught by Naylor, is deemed an obvious modification for better wire management as suggested by Morley.  
  
    **Conclusion:**  
    The examiner’s rejection under U.S.C 103 (Obviousness) or U.S.C 102 (Lack of Novelty)[Depending on which examiner claims] may be justified as the combination of features in the foundational claim appears to be an obvious modification of the Naylor reference, with the interconnection feature suggested by Morley.  
    """  
    return few_shot_example, text_a  
  
def generate_formatting_rules():  
    """Generate formatting rules for the analysis."""  
    return """  
    IMPORTANT FORMATTING RULES:  
    Numbering and Formatting: Use bullet points (•) instead of numbers when listing items.  
    Do not include markdown formatting in your response, except for bolding headings and sub-headings, and underlining as specified.  
    Bold all headings and sub-headings and the point headings for clarity.  
    Underline new language in the 'Proposed Amended Language' by enclosing it within '<u>' and '</u>' tags.  
    Provide detailed explanations and cite the sources correctly.  
    Propose amendments for all key features in foundational claim.  
    Do NOT include "N/A" anywhere and enclose words within asterisks(**)  
    Avoid one-line explanations; provide thorough and detailed analysis  
    Maintain concise formatting without extra line spacing  
    Add a conclusion after proposing amendments.  
    The provided examples are for structural guidance only; do not replicate them verbatim.  
    Avoid using words like "only" that may downplay the content  
    In the few shot examples, the text enclosed within **...** should be bold.  
    """  
  
def generate_analysis_prompt(extracted_details, foundational_claim, figure_analysis):  
    """Generate the user prompt for analyzing the filed application."""  
    formatting_rules = generate_formatting_rules()  
    return f"""  
    Analyze the filed application based on the foundational claim:  
    {json.dumps(foundational_claim, indent=2)}  
    and the figure analysis results:  
    {json.dumps(figure_analysis, indent=2)}  
    and the application as filed details:  
    {extracted_details}  
      
    {formatting_rules}  
      
    Key Features of Independent Claim with Number:  
    Extract and list the key features of the foundational claim. Ensure to include structural details, functional aspects, and any specific configurations mentioned in the claim.  
  
    Key Features of Cited Reference:  
    Extract and list the key features of the cited reference (also include where it is located in the cited text such as paragraph or figure). Highlight any similarities or differences in structure, function, and configuration compared to the foundational claim.  
  
    Examiner’s Analysis:  
    Describe the examiner’s analysis and the basis for rejection. Summarize how the examiner interprets the cited reference in relation to the foundational claim. Identify whether the rejection is based on U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness).  
  
    Novelty Analysis (U.S.C 102 - Lack of Novelty):  
    Compare the foundational claim with the cited reference to determine if the claim lacks novelty. Identify if all elements of the foundational claim are disclosed in the cited reference. Provide a detailed side-by-side comparison of each element.  
  
    Non-Obviousness Analysis (U.S.C 103 - Obviousness):  
    Analyze whether the foundational claim is obvious in light of the cited reference. Consider if the combination of features in the foundational claim would have been obvious to a person skilled in the art at the time of the invention. Discuss any differences that might contribute to non-obviousness.  
  
    Conclusion:  
    Provide a conclusion on whether the examiner’s rejection under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) may be justified or not. Summarize the key points that support or refute the examiner’s rejection.  
  
    Potential Areas for Distinction:  
    Identify areas where the foundational claim can be distinguished from the cited reference. Focus on unique structural features, specific materials, configurations, or functions not disclosed in the cited reference.  
  
    Proposed Amendments and Arguments:  
    For each key feature point in the foundational claim, propose specific amendments separately. NOTE: for all the points in the foundational claim, it is mandatory to propose amendments.  
    Present original and proposed versions, highlighting new features, specific materials, or configurations. **Underline** the new language proposed by enclosing it within '<u>' and '</u>' tags.  
      
    Format for Each Amendment:  
    Amendment [Number]: [Feature Title]  
    Original Claim Language:  
    "[Insert the exact original feature description from the foundational claim.]"  
    Proposed Amended Language:  
    "[Insert the enhanced feature description, incorporating new details, specific materials, or configurations. **Underline** the new language proposed by enclosing it within '<u>' and '</u>' tags.]"  
    Derivation of Amendment:  
    Source Reference: Cite specific sections, paragraphs, figures, or embodiments from the application that support the amendment. Example: "Derived from Paragraph [0123] and Figure 5 of the application."  
    Reasoning: Explain why the amendment was made, detailing how it enhances specificity, overcomes prior art, or adds technical advantages. Highlight any differences from the cited references. Emphasize any technical advantages or improvements introduced by the amendments.  
    """  
  
def format_analysis_output(response_content):  
    """Format the analysis output by removing code block markers and parsing JSON if necessary."""  
    # Remove markdown code block markers if present  
    if response_content.startswith("```json"):  
        response_content = response_content[7:-3].strip()  
    elif response_content.startswith("```"):  
        response_content = response_content[3:-3].strip()  
  
    # Attempt to parse the response as JSON  
    try:  
        return json.loads(response_content)  
    except json.JSONDecodeError:  
        print(response_content)  # Print the raw response if JSON decoding fails  
        return response_content  # Return the raw response as fallback  
  
def analyze_filed_application(extracted_details, foundational_claim, figure_analysis, domain, expertise, style):  
    """Analyze the filed application based on the foundational claim, figure analysis, and application details."""  
    content = generate_content(domain, expertise, style)  
    prompt = generate_analysis_prompt(extracted_details, foundational_claim, figure_analysis)  
    few_shot_example, text_a = generate_few_shot_examples()  
  
    messages = [  
        {"role": "system", "content": content},  
        {"role": "user", "content": prompt},  
        {"role": "user", "content": few_shot_example},  
        {"role": "user", "content": text_a},  
    ]  
  
    base_delay = 1  
    max_delay = 32  
    max_attempts = 5  
  
    for attempt in range(max_attempts):  
        try:  
            response = client.chat.completions.create(  
                model="GPT-4-Omni", messages=messages, temperature=0.2  
            )  
            analysis_output = response.choices[0].message.content.strip()  
            return format_analysis_output(analysis_output)  
  
        except Exception as e:  
            if attempt == max_attempts - 1:  
                logging.error(f"Max attempts reached. Error: {e}")  
                return None  
  
            delay = min(max_delay, base_delay * (2 ** attempt))  
            jitter = random.uniform(0, delay)  
            logging.warning(f"Retrying in {jitter:.2f} seconds (attempt {attempt + 1}) due to error: {e}")  
            time.sleep(jitter) 
  
  
def analyze_modified_application(cited_references_text, foundational_claim, figure_analysis, modified_application_details, domain, expertise, style): 
    content = f"""
   You are now assuming the role of a deeply specialized expert in {domain} as well as a comprehensive understanding of patent law specific to the mentioned domain. Your expertise includes:

    1. {domain}
    2. Patent Law Proficiency: 
    a. Skilled in interpreting and evaluating patent claims, classifications, and legal terminologies.
    b. Knowledgeable about the structure and requirements of patent applications.
    c. Expertise in comparing similar documents for patent claims under sections U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).

    3. {expertise}
    4. Capability to Propose Amendments:
    a. Experienced in responding to examiners’ assertions or rejections of claims.
    b. Skilled in proposing suitable amendments to patent claims to address rejections under U.S.C 102 (novelty) and U.S.C 103 (non-obviousness).
    c. Proficient in articulating and justifying amendments to ensure compliance with patentability requirements.

    Adopt a {style} suitable for analyzing patent applications in the given domain and subject matter. Your analysis should include:

    a. A thorough evaluation of the technical details and functionalities described in the patent application.
    b. An assessment of the clarity and precision of the technical descriptions and diagrams.
    c. An analysis of the novelty (under U.S.C 102) and non-obviousness (under U.S.C 103) of the subject matter by comparing it with similar existing documents.
    d. Feedback on the strengths and potential areas for improvement in the document.
    e. A determination of whether the invention meets the criteria for patentability under sections U.S.C 102 and U.S.C 103.
    f. Proposals for suitable amendments to the claims in response to potential examiners’ assertions or rejections, ensuring the claims are robust and meet patentability standards.

    Using this expertise, experience, and educational background, analyze the provided patent application document with a focus on its technical accuracy, clarity, adherence to patent application standards, novelty, non-obviousness, and overall feasibility.
    """ 
    prompt = f"""  
    Analyze the modified application based on the foundational claim:{json.dumps(foundational_claim, indent=2)}and the figure analysis results:{json.dumps(figure_analysis, indent=2)}and the modified application details:{json.dumps(modified_application_details, indent=2)}and the cited references:{json.dumps(cited_references_text, indent=2)}  
Assess whether the examiner's rejection of the application under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified by comparing it with the cited references text.
IMPORTANT FORMATTING RULES:
Numbering and Formatting:
Use bullet points (•) instead of numbers when listing items.
Do not include markdown formatting in your response.
Bolden only the headings.
Make your explanations lengthy and cite the sources correctly.
Give amendments for all key features in foundational claim.
Do NOT put N/A anywhere and enclose words within asterisks(**)
Key Features of Foundational Claim:
Extract and list the key features of the foundational claim.
Ensure to include structural details, functional aspects, and any specific configurations mentioned in the claim.
:Key Features of Cited Reference:
Extract and list the key features of the cited reference.(also include where it is located in the cited text)
Highlight any similarities or differences in structure, function, and configuration compared to the foundational claim.

Examiner’s Analysis:
Describe the examiner’s analysis and the basis for rejection.
Summarize how the examiner interprets the cited reference in relation to the foundational claim.
Identify whether the rejection is based on U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness).

Novelty Analysis (U.S.C 102 - Lack of Novelty):
Compare the foundational claim with the cited reference to determine if the claim lacks novelty.
Identify if all elements of the foundational claim are disclosed in the cited reference.
Provide a detailed side-by-side comparison of each element.

Non-Obviousness Analysis (U.S.C 103 - Obviousness):
Analyze whether the foundational claim is obvious in light of the cited reference.
Consider if the combination of features in the foundational claim would have been obvious to a person skilled in the art at the time of the invention.
Discuss any differences that might contribute to non-obviousness.

Conclusion:
Provide a conclusion on whether the examiner’s rejection under U.S.C 102 (Lack of Novelty) or U.S.C 103 (Obviousness) is justified.
Summarize the key points that support or refute the examiner’s rejection.

Potential Areas for Distinction:
Identify areas where the foundational claim can be distinguished from the cited reference.
Focus on unique structural features, specific materials, configurations, or functions not disclosed in the cited reference.

Proposed Amendments and Arguments:
For each key feature point in the foundational claim, propose specific amendments separately. NOTE: for all the points in the foundational claim, it is mandatory to propose amendments.
Present original and proposed versions, highlighting new features, specific materials, or configurations.
Amendment [Number]: [Feature]
Original: "[Original feature description...]"
Proposed: "[Enhanced feature description with new details, specific materials, or configurations...]"
Provide arguments supporting novelty and non-obviousness over the cited reference.
Emphasize any technical advantages or improvements introduced by the amendments.
IMPORTANT NOTE WHILE PROPOSING ARGUMENTS:
'''\Guidance for Proposing Amendments and Arguments:
When proposing amendments:

Be Specific: Clearly identify which feature you are amending and provide detailed enhancements.
Highlight Novel Elements: Emphasize new details such as specific materials, unique configurations, or innovative steps that are not present in the cited reference.
Refer to Sources: Cite sections of the application or figures from which the amendments and supporting arguments are drawn to reinforce their basis.
Maintain Claim Integrity: Ensure that the proposed amendments do not alter the fundamental essence of the original claim but enhance its patentability.
When crafting arguments to the examiner:
Address Rejection Points: Directly counter the examiner's reasons for rejection by highlighting differences between the amended claim and the cited reference.
Emphasize Novelty and Non-Obviousness: Explain why the amended features are new and not obvious, providing clear distinctions from the prior art.
Use Supporting Evidence: Reference specific examples, embodiments, or descriptions in the application that support your arguments.
Be Persuasive: Articulate the advantages and unique aspects of the invention that merit patent protection.\'''

Identify Limitations in Current Claims:
Identify any limitations or weaknesses in the current claims.
Propose specific language or structural changes to address these limitations.
Ensure that the proposed changes do not alter the original intent of the claims.

Propose New Arguments or Amendments:
Suggest additional arguments or amendments to further distinguish the foundational claim from the cited prior art.
Include multiple amendments for thorough differentiation.
Ensure that the original intent of the claims is maintained while improving clarity and scope.
    """  
      
    messages = [  
        {  
            "role": "system",  
            "content": content,

        },  
        {  
            "role": "user",  
            "content": prompt,  
        },  
    ]  
      
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni", messages=messages, temperature=0.6  
        )  
        analysis_output = response.choices[0].message.content.strip()  
          
        if analysis_output.startswith("```json"):  
            analysis_output = analysis_output[7:-3].strip()  
        elif analysis_output.startswith("```"):  
            analysis_output = analysis_output[3:-3].strip()  
          
        try:  
            return json.loads(analysis_output)  
        except json.JSONDecodeError:  
            return analysis_output  
    except Exception as e:  
        print(f"Error during modified application analysis: {e}")  
        return None  
  
  
def save_analysis_to_word(analysis_output):
    if not analysis_output:
        print("Analysis data is missing or empty.")
        return None

    # Create a new Word document
    doc = docx.Document()

    # Add a header with "Privileged & Confidential"
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "PRIVILEGED AND CONFIDENTIAL"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a heading for the document
    doc.add_heading("Filed Application Analysis Results", level=1)

    # Split the analysis output into lines
    lines = analysis_output.split("\n")
    for line in lines:
        line = line.strip()
        
        # Handle different levels of headings
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.", line):
            doc.add_paragraph(line, style="List Number")
        else:
            paragraph = doc.add_paragraph()
            
            # Regex to match bold text with optional trailing colon
            parts = re.split(r'(\*\*[^*]+\*\*:?|<u>.*?</u>)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**:"):
                    # Handle bold text with trailing colon
                    bold_text = part[2:-3]  # Remove ** and colon
                    run = paragraph.add_run(bold_text)
                    run.bold = True
                    paragraph.add_run(":")  # Add colon separately
                elif part.startswith("**") and part.endswith("**"):
                    # Handle bold text without trailing colon
                    run = paragraph.add_run(part[2:-2])  # Remove **
                    run.bold = True
                elif part.startswith("<u>") and part.endswith("</u>"):
                    # Handle underlined text
                    run = paragraph.add_run(part[3:-4])  # Remove <u> and </u>
                    run.underline = True
                else:
                    # Handle regular text
                    paragraph.add_run(part)

    # Save the document to a BytesIO buffer instead of writing to disk
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# Initialize session state variables  
session_vars = [  
    'conflict_results', 'foundational_claim', 'figure_analysis', 'filed_application_analysis',  
    'cited_documents', 'pending_claims_analysis', 'pending_claims_available', 'domain', 'expertise',  
    'style', 'filed_application_name'  
]  
  
for var in session_vars:  
    if var not in st.session_state:  
        st.session_state[var] = None  
  
st.session_state['pending_claims_available'] = st.session_state.get('pending_claims_available', "No")  

# Function to create aligned uploader and button  
def create_uploader_and_button(label_button, key):  
    col1, col2 = st.columns([4, 1])  # Adjust the column widths as needed  
    with col1:  
        uploaded_file = st.file_uploader("", type=["pdf", "docx"], key=key)  # Empty string for no label  
    with col2:  
        st.markdown("<br>", unsafe_allow_html=True)  # Add some space with HTML  
        button_clicked = st.button(label_button)  
    return uploaded_file, button_clicked  
  
def convert_docx_to_pdf(docx_path, pdf_path):  
    """Convert a DOCX file to PDF using docx2pdf."""  
    try:  
        convert(docx_path, pdf_path)  
        return pdf_path  
    except Exception as e:  
        st.error(f"Failed to convert DOCX to PDF: {e}")  
        return None  
  
def extract_text_from_pdf(uploaded_pdf_path):  
    """Extract text from a PDF file using Azure Form Recognizer Document Intelligence."""  
    try:  
        # Initialize DocumentAnalysisClient  
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key),  
        )  
  
        # Read the file content  
        with open(uploaded_pdf_path, "rb") as f:  
            file_content = f.read()  
  
        # Use the prebuilt-document model to analyze the document  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=file_content  
        )  
  
        # Get the result of the analysis  
        result = poller.result()  
  
        # Extract the text from the result  
        text = ""  
        for page in result.pages:  
            for line in page.lines:  
                text += line.content + "\n"  
  
        return text  
  
    except HttpResponseError as e:  
        st.error(f"Failed to analyze the document: {e.message}")  
        return None  
  
    except Exception as e:  
        st.error(f"An unexpected error occurred: {e}")  
        return None  
  
# Function to convert DOCX to PDF  
def convert_word_to_pdf(input_file, output_file):  
    try:  
        pypandoc.convert_file(input_file, 'pdf', outputfile=output_file, extra_args=['--pdf-engine=pdflatex'])  
        return output_file  
    except Exception as e:  
        st.error(f"Error converting file: {e}")  
        return None  
  
# Function to merge multiple PDFs  
def merge_pdfs(pdf_list, output_file):  
    merger = PdfMerger()  
    for pdf in pdf_list:  
        merger.append(pdf)  
    merger.write(output_file)  
    merger.close()  
    return output_file  
def validate_office_action(uploaded_file):  
    if not uploaded_file:  
        st.error("No file uploaded.")  
        return False, None, None  
  
    try:  
        file_content = uploaded_file.read()  
  
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key),  
        )  
  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=file_content  
        )  
  
        result = poller.result()  
  
        application_number = None  
        conflict_keyword = None  
        summary_found = False  
  
        for page in result.pages:  
            for line in page.lines:  
                content = line.content.lower()  
  
                if "application no" in content or "control number" in content:  
                    application_number = line.content.split()[-1]  
  
                if "office action summary" in content:  
                    summary_found = True  
  
                if "rejected" in content and "102(a)(1)" in content:  
                    match = re.search(r"by (\w+)", line.content)  
                    if match:  
                        conflict_keyword = match.group(1)  
  
        if application_number and summary_found:  
            return True, application_number, conflict_keyword  
  
        st.error("The uploaded document is not a valid Office Action.")  
        return False, None, None  
  
    except HttpResponseError as e:  
        st.error(f"Failed to analyze the document: {e.message}")  
        return False, None, None 
    
def validate_application_as_filed(uploaded_file, expected_application_number):  
    if not uploaded_file:  
        st.error("No file uploaded.")  
        return False  
  
    if expected_application_number is None:  
        st.error("Application number is not set.")  
        return False  
  
    try:  
        # Read the file content  
        with open(uploaded_file, "rb") as f:  
            file_content = f.read()  
          
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key),  
        )  
  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=file_content  
        )  
  
        result = poller.result()  
  
        for page in result.pages:  
            for line in page.lines:  
                if expected_application_number in line.content:  
                    st.success("Application as Filed validated successfully!")  
                    return True  
  
        st.error(f"The document does not contain the expected application number: {expected_application_number}.")  
        return False  
  
    except HttpResponseError as e:  
        st.error(f"Failed to analyze the document: {e.message}")  
        return False  
def match_document_name_or_pub_number(file_name, cited_docs):  
    # Normalize and prepare regex pattern for matching  
    file_name = file_name.lower()  
    file_name_pattern = re.sub(r'[^a-z0-9]', '', file_name)  # Remove non-alphanumeric characters for comparison  
  
    for cited_doc in cited_docs:  
        # Normalize cited document name for comparison  
        cited_doc_name = cited_doc.lower()  
        cited_doc_pattern = re.sub(r'[^a-z0-9]', '', cited_doc_name)  
  
        # Check if file name matches the cited document  
        if re.search(file_name_pattern, cited_doc_pattern):  
            return True  
    return False  
def extract_text_from_pdfs(file_path):  
    try:  
        document_analysis_client = DocumentAnalysisClient(  
            endpoint=form_recognizer_endpoint,  
            credential=AzureKeyCredential(form_recognizer_api_key),  
        )  
  
        with open(file_path, "rb") as f:  
            file_content = f.read()  
  
        poller = document_analysis_client.begin_analyze_document(  
            "prebuilt-document", document=file_content  
        )  
        result = poller.result()  
  
        text = ""  
        for page in result.pages[:2]:  
            for line in page.lines:  
                text += line.content + "\n"  
        return text  
    except HttpResponseError as e:  
        st.error(f"Failed to analyze the document: {e.message}")  
        return None  
    except Exception as e:  
        st.error(f"An unexpected error occurred: {e}")  
        return None  
  
def check_match_with_llm(text, cited_docs):  
    messages = [  
        {  
            "role": "system",  
            "content": "Identify if any part of the text approximately matches any item in the given array. Respond with 'Yes' or 'No'."  
        },  
        {  
            "role": "user",  
            "content": f"Text: {text}\nCited Documents: {cited_docs}"  
        }  
    ]  
  
    try:  
        response = client.chat.completions.create(  
            model="GPT-4-Omni",  
            messages=messages,  
            temperature=0.2  
        )  
          
        llm_response = response.choices[0].message.content.strip()  
        st.write(f"LLM Response: {llm_response}")  
  
        match_found = llm_response.lower() == "yes"  
        return match_found  
    except Exception as e:  
        st.error(f"Error during LLM check: {e}")  
        return False  
  
  
# Ensure session state is initialized  
if 'conflict_results' not in st.session_state:  
    st.session_state.conflict_results = None  
if 'foundational_claim' not in st.session_state:  
    st.session_state.foundational_claim = None  
if 'figure_analysis' not in st.session_state:  
    st.session_state.figure_analysis = None  
if 'filed_application_analysis' not in st.session_state:  
    st.session_state.filed_application_analysis = None  
if 'cited_documents' not in st.session_state:  
    st.session_state.cited_documents = None  
if 'pending_claims_analysis' not in st.session_state:  
    st.session_state.pending_claims_analysis = None  
if 'pending_claims_available' not in st.session_state:  
    st.session_state.pending_claims_available = "No"  
if 'domain' not in st.session_state:  
    st.session_state.domain = None  
if 'expertise' not in st.session_state:  
    st.session_state.expertise = None  
if 'style' not in st.session_state:  
    st.session_state.style = None  
if 'filed_application_name' not in st.session_state:  
    st.session_state.filed_application_name = None  
if 'application_number' not in st.session_state:  
    st.session_state.application_number = None  
  
# Display the logo and title  
st.image("AFS Innovation Logo.png", width=200)  
st.title("Patent Analyzer")  
  
# Step 1: Upload Examiner Document and Check Conflicts  
with st.expander("Step 1: Office Action", expanded=True):  
    st.write("### Upload the Examiner Document and Check for Conflicts")  
    uploaded_examiner_file = st.file_uploader("Upload Examiner Document", type=["pdf", "docx"])  
    conflicts_clicked = st.button("Check for Conflicts")  
  
    if conflicts_clicked:  
        if uploaded_examiner_file is None:  
            st.warning("Please upload the examiner document first.")  
        else:  
            is_valid, application_number, conflict_keyword = validate_office_action(uploaded_examiner_file)  
            if not is_valid:  
                st.error("Failed to process the uploaded file.")  
            else:  
                st.session_state.application_number = application_number  
  
                # Reset file pointer to read again  
                uploaded_examiner_file.seek(0)  
                temp_file_path = "temp_examiner.pdf" if uploaded_examiner_file.type == "application/pdf" else "temp_examiner.docx"  
                with open(temp_file_path, "wb") as f:  
                    f.write(uploaded_examiner_file.read())  
  
                # Upload the input file to Azure Blob Storage  
                folder_name = st.session_state.application_number  # Use application number as folder name  
                upload_file_to_blob(temp_file_path, uploaded_examiner_file.name, folder_name)  
  
                if uploaded_examiner_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  
                    temp_pdf_path = "temp_examiner_converted.pdf"  
                    pdf_path = convert_docx_to_pdf(temp_file_path, temp_pdf_path)  
                    if not pdf_path:  
                        st.error("Failed to convert DOCX to PDF.")  
                    else:  
                        temp_file_path = pdf_path  # Update the path to the converted PDF  
  
                if os.path.exists(temp_file_path):  
                    extracted_examiner_text = extract_and_process_text_from_pdf(temp_file_path)  
                    if not extracted_examiner_text:  
                        st.error("Failed to extract text from the examiner document.")  
                    else:  
                        # Process the extracted text  
                        processed_examiner_text = "\n".join(extracted_examiner_text)  # Join pages for further processing  
  
                        domain, expertise, style = determine_domain_expertise(processed_examiner_text)  
                        if not (domain and expertise and style):  
                            st.error("Failed to determine domain expertise.")  
                        else:  
                            st.session_state.domain = domain  
                            st.session_state.expertise = expertise  
                            st.session_state.style = style  
  
                            conflict_results_raw = check_for_conflicts(processed_examiner_text, domain, expertise, style)  
                            if not conflict_results_raw:  
                                     st.error("Failed to check for conflicts.")  
                            else:  
                               st.session_state.conflict_results = conflict_results_raw  
                               st.session_state.foundational_claim = conflict_results_raw.get("foundational_claim")  
                               st.session_state.cited_documents = conflict_results_raw.get("documents_referenced")  
                               st.success("Conflicts checked successfully!")  
  
                    os.remove(temp_file_path) 
                    if uploaded_examiner_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" and os.path.exists(temp_pdf_path):  
                        os.remove(temp_pdf_path)  
# Check if cited documents exist  
if st.session_state.get("cited_documents") is not None:  
    st.write("### Cited Documents Referenced:")  
    cited_docs_df = pd.DataFrame({  
        "No.": range(1, len(st.session_state.cited_documents) + 1),  
        "Document Name": st.session_state.cited_documents  
    })  
    with st.container():  
        st.dataframe(cited_docs_df.set_index("No."), use_container_width=True)  
  
# Step 2: Upload Referenced Document and Analyze Figures  
if st.session_state.get("conflict_results") is not None:  
    with st.expander("Step 2: Referenced Documents", expanded=True):  
        st.write("### Upload the Referenced Documents and Analyze Figures")  
        uploaded_ref_files = st.file_uploader("", type="pdf", key="referenced", accept_multiple_files=True)  
        analyze_figures_clicked = st.button("Analyze Figures and Cited Text")  
  
        if analyze_figures_clicked:  
            if uploaded_ref_files:  
                ref_texts = []  
                cited_docs = st.session_state.cited_documents  
                continue_processing = False  
  
                for uploaded_ref_file in uploaded_ref_files:  
                    temp_file_path = f"temp_{uploaded_ref_file.name}"  
                    with open(temp_file_path, "wb") as f:  
                        f.write(uploaded_ref_file.read())  
  
                    # Upload the input file to Azure Blob Storage  
                    folder_name = st.session_state.application_number  
                    upload_file_to_blob(temp_file_path, uploaded_ref_file.name, folder_name)  
  
                    extracted_ref_text = extract_and_process_text_from_pdf(temp_file_path)  
                    if extracted_ref_text:  
                        match_found = check_match_with_llm("\n".join(extracted_ref_text), cited_docs)  
                        if match_found:  
                            processed_ref_text = "\n".join(extracted_ref_text)  
                            ref_texts.append(processed_ref_text)  
                            continue_processing = True  # Set flag to continue  
                        else:  
                            st.warning(f"No match for cited documents was found in {uploaded_ref_file.name}.")  
  
                    os.remove(temp_file_path)  
  
                if continue_processing:  
                    # Perform figure analysis if a match was found  
                    figure_analysis_results = extract_figures_and_text(  
                        st.session_state.conflict_results, " ".join(ref_texts),  
                        st.session_state.domain, st.session_state.expertise, st.session_state.style  
                    )  
  
                    if figure_analysis_results:  
                        st.session_state.figure_analysis = figure_analysis_results  
                        st.success("Figure analysis completed successfully!")  
                    else:  
                        st.error("Failed to analyze figures and cited text.")  
            else:  
                st.warning("Please upload the referenced documents first.")  
  
# Step 3: Ask if the Application is Published  
if st.session_state.get("figure_analysis") is not None:  
    with st.expander("Step 3: Application as Filed", expanded=True):  
        st.write("### Is the Application Published?")  
        is_published = st.radio("Select an option:", ("Yes", "No"))  
  
        if is_published == "No":  
            st.write("### Upload the DOCX and PDF to Combine and Analyze")  
            word_file = st.file_uploader("Upload Word document", type=["docx"])  
            pdf_file = st.file_uploader("Upload PDF document", type=["pdf"])  
            combine_and_proceed_clicked = st.button("Combine and Proceed")  
  
            if combine_and_proceed_clicked:  
                if word_file and pdf_file:  
                    with tempfile.TemporaryDirectory() as tmpdirname:  
                        word_path = os.path.join(tmpdirname, word_file.name)  
                        pdf_path = os.path.join(tmpdirname, pdf_file.name)  
  
                        with open(word_path, "wb") as f:  
                            f.write(word_file.getbuffer())  
                        with open(pdf_path, "wb") as f:  
                            f.write(pdf_file.getbuffer())  
  
                        # Upload the input files to Azure Blob Storage  
                        folder_name = st.session_state.application_number  
                        upload_file_to_blob(word_path, word_file.name, folder_name)  
                        upload_file_to_blob(pdf_path, pdf_file.name, folder_name)  
  
                        output_pdf_file = os.path.join(tmpdirname, "combined_document.pdf")  
                        with st.spinner("Converting Word to PDF..."):  
                            converted_pdf = convert_docx_to_pdf(word_path, os.path.join(tmpdirname, "converted.pdf"))  
                        if converted_pdf:  
                            with st.spinner("Merging PDFs..."):  
                                merged_pdf = merge_pdfs([converted_pdf, pdf_path], output_pdf_file)  
                            st.success("DOCX and PDF have been successfully combined!")  
  
                            # Upload the combined PDF to Azure Blob Storage  
                            upload_file_to_blob(output_pdf_file, "combined_document.pdf", folder_name)  
  
                            with open(output_pdf_file, "rb") as f:  
                                st.download_button(  
                                    label="Download Combined PDF",  
                                    data=f,  
                                    file_name="combined_document.pdf",  
                                    mime="application/pdf"  
                                )  
  
                            st.session_state.filed_application_name = pdf_file.name  
                            # Proceed with Step 3 as the combined PDF is ready  
                            extracted_filed_app_text = extract_and_process_text_from_pdf(output_pdf_file)  
  
                            if extracted_filed_app_text:  
                                processed_filed_app_text = "\n".join(extracted_filed_app_text)  
  
                                filed_app_details = extract_details_from_filed_application(  
                                    processed_filed_app_text,  
                                    st.session_state.foundational_claim,  
                                    st.session_state.domain,  
                                    st.session_state.expertise,  
                                    st.session_state.style  
                                )  
                                if filed_app_details:  
                                    filed_app_details_json = json.dumps(filed_app_details, indent=2)  
                                    st.session_state.filed_application_analysis = filed_app_details_json  
  
                                    analysis_results = analyze_filed_application(  
                                        filed_app_details_json,  
                                        st.session_state.foundational_claim,  
                                        st.session_state.figure_analysis,  
                                        st.session_state.domain,  
                                        st.session_state.expertise,  
                                        st.session_state.style  
                                    )  
                                    if analysis_results:  
                                        st.session_state.filed_application_analysis = analysis_results  
                                        st.success("Filed application analysis completed successfully!")  
                                        docx_buffer = save_analysis_to_word(analysis_results)  
                                        if docx_buffer:  
                                            filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
                                            st.download_button(  
                                                label="Download Analysis Results",  
                                                data=docx_buffer,  
                                                file_name=f"{filed_application_name}_ANALYSIS.docx",  
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                                key="filed_application_download"  
                                            )  
                                    else:  
                                        st.error("Failed to analyze the filed application.")  
                                else:  
                                    st.error("Failed to analyze the filed application.")  
                            else:  
                                st.error("Failed to extract text from the filed application document.")  
                        else:  
                            st.error("Failed to convert Word to PDF.")  
                else:  
                    st.warning("Please upload both the DOCX and PDF files.")  
  
        elif is_published == "Yes":  
            uploaded_filed_app = st.file_uploader("Upload Filed Application", type=["pdf"])  
            analyze_filed_app_clicked = st.button("Analyze Filed Application")  
  
            if analyze_filed_app_clicked:  
                if uploaded_filed_app is not None:  
                    temp_file_path = "temp_filed.pdf"  
                    with open(temp_file_path, "wb") as f:  
                        f.write(uploaded_filed_app.read())  
  
                    # Upload the input file to Azure Blob Storage  
                    folder_name = st.session_state.application_number  
                    upload_file_to_blob(temp_file_path, uploaded_filed_app.name, folder_name)  
  
                    # Validate the uploaded filed application  
                    extracted_filed_app_text = extract_and_process_text_from_pdf(temp_file_path)  
                    os.remove(temp_file_path)  
  
                    if extracted_filed_app_text:  
                        processed_filed_app_text = "\n".join(extracted_filed_app_text)  
                        st.session_state.filed_application_name = uploaded_filed_app.name  
  
                        filed_app_details = extract_details_from_filed_application(  
                            processed_filed_app_text,  
                            st.session_state.foundational_claim,  
                            st.session_state.domain,  
                            st.session_state.expertise,  
                            st.session_state.style  
                        )  
                        if filed_app_details:  
                            filed_app_details_json = json.dumps(filed_app_details, indent=2)  
                            st.session_state.filed_application_analysis = filed_app_details_json  
  
                            analysis_results = analyze_filed_application(  
                                filed_app_details_json,  
                                st.session_state.foundational_claim,  
                                st.session_state.figure_analysis,  
                                st.session_state.domain,  
                                st.session_state.expertise,  
                                st.session_state.style  
                            )  
  
                            if analysis_results:  
                                st.session_state.filed_application_analysis = analysis_results  
                                st.success("Filed application analysis completed successfully!")  
                                docx_buffer = save_analysis_to_word(analysis_results)  
                                if docx_buffer:  
                                    filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
                                    st.download_button(  
                                        label="Download Analysis Results",  
                                        data=docx_buffer,  
                                        file_name=f"{filed_application_name}_ANALYSIS.docx",  
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                        key="filed_application_download"  
                                    )  
                            else:  
                                st.error("Failed to analyze the filed application.")  
                        else:  
                            st.error("Failed to analyze the filed application.")  
                    else:  
                        st.error("Failed to extract text from the filed application document.")  
                else:  
                    st.warning("Please upload the filed application first.")  
  
# STEP 4: Pending Claims Analysis  
if st.session_state.get("filed_application_analysis") is not None:  
    with st.expander("Step 4: Pending Claims", expanded=True):  
        st.write("### Do you have a Pending Claims Document to Analyze?")  
        st.session_state.pending_claims_available = st.radio(  
            "Select an option:",  
            ("Yes", "No"),  
            index=0 if st.session_state.pending_claims_available == "Yes" else 1,  
            key="pending_claims_radio"  
        )  
  
        if st.session_state.pending_claims_available == "Yes":  
            st.write("### Upload the Pending Claims Document and Analyze")  
            uploaded_pending_claims_file = st.file_uploader("Upload Pending Claims Document", type=["pdf", "docx"])  
            analyze_pending_claims_clicked = st.button("Analyze Pending Claims")  
  
            if analyze_pending_claims_clicked:  
                if uploaded_pending_claims_file is not None:  
                    with tempfile.TemporaryDirectory() as tmpdirname:  
                        file_path = os.path.join(tmpdirname, uploaded_pending_claims_file.name)  
                        with open(file_path, "wb") as f:  
                            f.write(uploaded_pending_claims_file.read())  
  
                        # Upload the input file to Azure Blob Storage  
                        folder_name = st.session_state.application_number  
                        upload_file_to_blob(file_path, uploaded_pending_claims_file.name, folder_name)  
  
                        # Validate the pending claims document  
                        extracted_pending_claims_text = None  
                        if uploaded_pending_claims_file.type == "application/pdf":  
                            extracted_pending_claims_text = extract_and_process_text_from_pdf(file_path)  
                        elif uploaded_pending_claims_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":  
                            extracted_pending_claims_text = extract_text_from_docx(file_path)  
  
                        if extracted_pending_claims_text:  
                            processed_pending_claims_text = "\n".join(extracted_pending_claims_text)  
  
                            modified_filed_application_results = extract_and_modify_filed_application(  
                                st.session_state.filed_application_analysis,  
                                processed_pending_claims_text,  
                                st.session_state.domain,  
                                st.session_state.expertise,  
                                st.session_state.style  
                            )  
  
                            if modified_filed_application_results:  
                                st.session_state.modified_filed_application_results = modified_filed_application_results  
                                st.success("Modified filed application analysis completed successfully!")  
  
                                pending_claims_analysis_results = analyze_modified_application(  
                                    processed_pending_claims_text,  
                                    st.session_state.foundational_claim,  
                                    st.session_state.figure_analysis,  
                                    modified_filed_application_results,  
                                    st.session_state.domain,  
                                    st.session_state.expertise,  
                                    st.session_state.style  
                                )  
  
                                if pending_claims_analysis_results:  
                                    st.session_state.pending_claims_analysis = pending_claims_analysis_results  
                                    st.success("Pending claims analysis completed successfully!")  
  
                                    docx_buffer = save_analysis_to_word(pending_claims_analysis_results)  
                                    if docx_buffer:  
                                        filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
                                        st.download_button(  
                                            label="Download Analysis Results",  
                                            data=docx_buffer,  
                                            file_name=f"{filed_application_name}_ANALYSIS.docx",  
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
                                            key="pending_claims_download"  
                                        )  
                                else:  
                                    st.error("Failed to analyze the pending claims.")  
                            else:  
                                st.error("Failed to modify the filed application based on pending claims.")  
                        else:  
                            st.error("Failed to extract text from the pending claims document.")  
                else:  
                    st.warning("Please upload the pending claims document first.")  
  
# Option to download results if there are no pending claims  
if st.session_state.get("filed_application_analysis") and st.session_state.pending_claims_analysis is None:  
    docx_buffer = save_analysis_to_word(st.session_state.filed_application_analysis)  
    if docx_buffer:  
        filed_application_name = st.session_state.filed_application_name.replace(" ", "_")  
        st.download_button(  
            label="Download Analysis Results",  
            data=docx_buffer,  
            file_name=f"{filed_application_name}_ANALYSIS.docx",  
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",  
            key="filed_application_final_download"  
        )  
